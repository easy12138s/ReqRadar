"""DOCX 文件加载器"""

import logging
from pathlib import Path

from reqradar.modules.loaders.base import DocumentLoader, LoadedDocument

logger = logging.getLogger("reqradar.loaders.docx")

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxLoader(DocumentLoader):
    def __init__(self, chunk_size: int = 300, chunk_overlap: int = 50):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def load(self, file_path: Path, **kwargs) -> list[LoadedDocument]:
        chunk_size = kwargs.get("chunk_size", self.chunk_size)
        chunk_overlap = kwargs.get("chunk_overlap", self.chunk_overlap)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        if not full_text.strip():
            logger.warning("No text extracted from DOCX: %s", file_path)
            return []

        chunks = _chunk_docx_text(full_text, chunk_size=chunk_size, overlap=chunk_overlap)

        return [
            LoadedDocument(
                content=chunk,
                source=str(file_path),
                format="docx",
                metadata={"title": file_path.stem, "chunk_index": i},
            )
            for i, chunk in enumerate(chunks)
        ]


def _chunk_docx_text(text: str, chunk_size: int = 300, overlap: int = 50) -> list[str]:
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if start > 0:
            chunk = f"[接上文]\n{chunk}"
        chunks.append(chunk)
        start = end - overlap
    return chunks
